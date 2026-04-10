import sys
import os
import tkinter as tk
from tkinter import messagebox, scrolledtext, filedialog
from docx import Document
import re
try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

# Translation table for stripping invisible formatting characters
# (soft hyphens, zero-width spaces, etc. that Word inserts for formatting)
_INVISIBLE_CHARS_TABLE = str.maketrans('', '', '\u00AD\u200B\u200C\u200D\u2060\uFEFF')

class WordTextReplacerSingle:
    def __init__(self, initial_file=None):
        self.file_paths = []
        self._text_cache = {}
        self._live_count_after_id = None
        if initial_file and os.path.exists(initial_file):
            self.file_paths.append(initial_file)
        
        self.root = tk.Tk()
        self.setup_gui()
        
        # Build initial text cache
        if self.file_paths:
            self._refresh_text_cache()
        
    def setup_gui(self):
        self.update_title()
        self.root.geometry("700x780")
        self.root.resizable(True, True)
        
        # Main frame
        main_frame = tk.Frame(self.root, padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Header frame with info icon
        header_frame = tk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 15))
        
        # Invisible spacer to push info icon to the right
        spacer = tk.Label(header_frame, text="")
        spacer.pack(side=tk.LEFT, expand=True)
        
        # Info icon with tooltip
        self.info_icon = tk.Label(header_frame, text="?", font=("Arial", 12, "bold"), 
                                 fg="#666666", cursor="hand2", 
                                 relief="solid", borderwidth=1, padx=4, pady=2)
        self.info_icon.pack(side=tk.RIGHT)
        
        # Bind hover events for tooltip
        self.info_icon.bind("<Enter>", self.show_shortcuts_tooltip)
        self.info_icon.bind("<Leave>", self.hide_shortcuts_tooltip)
        
        # File management section
        files_frame = tk.LabelFrame(main_frame, text="Selected files", 
                                   font=("Arial", 11, "bold"), padx=10, pady=10)
        files_frame.pack(fill=tk.X, pady=(0, 15))
        
        # File list with scrollbar
        list_frame = tk.Frame(files_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        self.file_listbox = tk.Listbox(list_frame, height=4, font=("Arial", 9), selectmode=tk.EXTENDED)
        scrollbar = tk.Scrollbar(list_frame, orient="vertical")
        self.file_listbox.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.file_listbox.yview)
        
        # Bind Delete key to remove files when listbox is focused
        self.file_listbox.bind('<Delete>', self.handle_delete_key)
        
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # File management buttons
        file_buttons_frame = tk.Frame(files_frame)
        file_buttons_frame.pack(fill=tk.X, pady=(10, 0))
        
        add_files_btn = tk.Button(file_buttons_frame, text="Add more files", 
                                 command=self.add_files,
                                 bg="#2196f3", fg="white", font=("Arial", 9))
        add_files_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        remove_file_btn = tk.Button(file_buttons_frame, text="Remove selected", 
                                   command=self.remove_selected_file,
                                   bg="#ff9800", fg="white", font=("Arial", 9))
        remove_file_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        clear_files_btn = tk.Button(file_buttons_frame, text="Clear all", 
                                   command=self.clear_all_files,
                                   bg="#f44336", fg="white", font=("Arial", 9))
        clear_files_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        # NEW: Add hyperlink check button here
        hyperlink_btn = tk.Button(file_buttons_frame, text="🔗 Hyperlink check", 
                                 command=self.check_hyperlinks,
                                 bg="#9c27b0", fg="white", font=("Arial", 9))
        hyperlink_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        # File count label
        self.file_count_label = tk.Label(file_buttons_frame, text="", 
                                        font=("Arial", 9), fg="gray")
        self.file_count_label.pack(side=tk.RIGHT)
        
        # Search text section
        search_header_frame = tk.Frame(main_frame)
        search_header_frame.pack(fill=tk.X, pady=(0, 5))
        
        search_label = tk.Label(search_header_frame, text="Search for:", 
                               font=("Arial", 11, "bold"))
        search_label.pack(side=tk.LEFT)
        
        paste_search_btn = tk.Button(search_header_frame, text="📋 Paste clipboard", 
                                    command=self.paste_to_search,
                                    bg="#e8f5e8", font=("Arial", 9))
        paste_search_btn.pack(side=tk.RIGHT)

        debug_btn = tk.Button(search_header_frame, text="🔍 nbsp check", 
                             command=self.debug_nbsp_conversion,
                             bg="#ffeb3b", font=("Arial", 9))
        debug_btn.pack(side=tk.RIGHT, padx=(10, 5))
        
        self.search_text = scrolledtext.ScrolledText(main_frame, height=6, width=70,
                                                    wrap=tk.WORD, font=("Arial", 10), undo=True)
        self.search_text.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        # Replace text section
        replace_header_frame = tk.Frame(main_frame)
        replace_header_frame.pack(fill=tk.X, pady=(0, 5))
        
        replace_label = tk.Label(replace_header_frame, text="Replace with:", 
                                font=("Arial", 11, "bold"))
        replace_label.pack(side=tk.LEFT)
        
        paste_replace_btn = tk.Button(replace_header_frame, text="📋 Paste clipboard", 
                                     command=self.paste_to_replace,
                                     bg="#e8f5e8", font=("Arial", 9))
        paste_replace_btn.pack(side=tk.RIGHT)
        
        self.replace_text = scrolledtext.ScrolledText(main_frame, height=6, width=70,
                                                     wrap=tk.WORD, font=("Arial", 10), undo=True)
        self.replace_text.pack(fill=tk.BOTH, expand=True, pady=(0, 20))
        
        # Options frame
        options_frame = tk.Frame(main_frame)
        options_frame.pack(fill=tk.X, pady=(0, 15))
        
        # First row of options
        options_row1 = tk.Frame(options_frame)
        options_row1.pack(fill=tk.X)
        
        self.create_backup_var = tk.BooleanVar(value=False)
        backup_checkbox = tk.Checkbutton(options_row1, text="Create backup files (.backup)", 
                                        variable=self.create_backup_var, font=("Arial", 10))
        backup_checkbox.pack(side=tk.LEFT)
        
        # Second row of options
        options_row2 = tk.Frame(options_frame)
        options_row2.pack(fill=tk.X, pady=(5, 0))
        
        self.case_sensitive_var = tk.BooleanVar(value=False)
        case_checkbox = tk.Checkbutton(options_row2, text="Case sensitive search", 
                                      variable=self.case_sensitive_var, font=("Arial", 10))
        case_checkbox.pack(side=tk.LEFT)
        
        # Third row of options
        options_row3 = tk.Frame(options_frame)
        options_row3.pack(fill=tk.X, pady=(5, 0))
        
        self.regex_var = tk.BooleanVar(value=False)
        regex_checkbox = tk.Checkbutton(options_row3, text="Regex (Standard Replace only)", 
                                       variable=self.regex_var, font=("Arial", 10))
        regex_checkbox.pack(side=tk.LEFT)
        
        # Fourth row of options
        options_row4 = tk.Frame(options_frame)
        options_row4.pack(fill=tk.X, pady=(5, 0))
        
        self.whole_word_var = tk.BooleanVar(value=False)
        whole_word_checkbox = tk.Checkbutton(options_row4, text="Whole word match", 
                                            variable=self.whole_word_var, font=("Arial", 10))
        whole_word_checkbox.pack(side=tk.LEFT)
        
        # Buttons frame
        button_frame = tk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Preview button
        preview_btn = tk.Button(button_frame, text="Preview", 
                               command=self.preview_changes,
                               bg="#e3f2fd", font=("Arial", 10))
        preview_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Standard Replace button
        self.replace_btn = tk.Button(button_frame, text="Standard Replace", 
                                    command=self.replace_text_in_documents,
                                    bg="#4caf50", fg="white", font=("Arial", 10, "bold"))
        self.replace_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Advanced Replace button
        advanced_replace_btn = tk.Button(button_frame, text="Advanced Replace", 
                                        command=self.advanced_replace_with_vba,
                                        bg="#ff9800", fg="white", font=("Arial", 10, "bold"))
        advanced_replace_btn.pack(side=tk.LEFT, padx=(0, 10))
                
        # Cancel button
        cancel_btn = tk.Button(button_frame, text="Close", 
                              command=self.root.destroy,
                              bg="#333333", fg="white", font=("Arial", 10))
        cancel_btn.pack(side=tk.RIGHT)
        
        # Progress bar frame (initially hidden)
        self.progress_frame = tk.Frame(main_frame)
        self.progress_label = tk.Label(self.progress_frame, text="", font=("Arial", 9))
        self.progress_label.pack(anchor="w")
        
        # Progress percentage label
        self.progress_percent_label = tk.Label(self.progress_frame, text="", font=("Arial", 9), fg="blue")
        self.progress_percent_label.pack(anchor="w")
        
        # Status label
        self.status_label = tk.Label(main_frame, text="Ready", fg="green", font=("Arial", 9))
        self.status_label.pack(anchor="w")
        
        # Match counter label (live counting)
        self.match_counter_label = tk.Label(main_frame, text="", fg="#555555", font=("Arial", 9, "italic"))
        self.match_counter_label.pack(anchor="w")
        
        # Setup keyboard bindings
        self.setup_keyboard_bindings()
        
        # Update the file list display
        self.update_file_list()
        
        # Bind search text changes for live counter
        self.search_text.bind('<KeyRelease>', self._on_search_key_release)
        
        # Also trigger live count when options change
        self.case_sensitive_var.trace_add('write', lambda *_: self._schedule_live_count())
        self.regex_var.trace_add('write', lambda *_: self._schedule_live_count())
        self.whole_word_var.trace_add('write', lambda *_: self._schedule_live_count())
        
        # Focus on search field
        self.search_text.focus()

    # NEW: NBSP Processing Methods
    def preprocess_text_with_nbsp(self, text):
        """Convert _nbsp_ placeholders to actual non-breaking spaces"""
        return text.replace('_nbsp_', '\u00A0')

    def _strip_invisible_chars(self, text):
        """Remove invisible formatting characters (soft hyphens, zero-width spaces, etc.)"""
        return text.translate(_INVISIBLE_CHARS_TABLE)

    def get_processed_search_text(self):
        """Get search text with _nbsp_ converted and invisible chars stripped"""
        raw_text = self.search_text.get("1.0", tk.END).strip()
        return self._strip_invisible_chars(self.preprocess_text_with_nbsp(raw_text))

    def get_processed_replace_text(self):
        """Get replace text with _nbsp_ converted to actual non-breaking spaces"""
        raw_text = self.replace_text.get("1.0", tk.END).strip()
        return self.preprocess_text_with_nbsp(raw_text)

    def debug_nbsp_conversion(self):
        """Debug method to show _nbsp_ conversion"""
        raw_search = self.search_text.get("1.0", tk.END).strip()
        processed_search = self.get_processed_search_text()
        raw_replace = self.replace_text.get("1.0", tk.END).strip()
        processed_replace = self.get_processed_replace_text()
        
        # Define the NBSP character outside the f-string
        nbsp_char = '\u00A0'
        
        debug_info = "nbsp found:\n\n"
        debug_info += f"Raw search text: '{raw_search}'\n"
        debug_info += f"Processed search text: '{processed_search}'\n"
        debug_info += f"nbsp in search text: {processed_search.count(nbsp_char)}\n\n"
        debug_info += f"Raw replace text: '{raw_replace}'\n"
        debug_info += f"Processed replace text: '{processed_replace}'\n"
        debug_info += f"nbsp in replace text: {processed_replace.count(nbsp_char)}\n\n"
        
        messagebox.showinfo("nbsp check", debug_info)

    # ── Live Match Counter ──
    def _on_search_key_release(self, event=None):
        self._schedule_live_count()

    def _schedule_live_count(self):
        if self._live_count_after_id is not None:
            self.root.after_cancel(self._live_count_after_id)
        self._live_count_after_id = self.root.after(400, self._do_live_count)

    def _refresh_text_cache(self):
        """Re-read document text for all current files"""
        self._text_cache.clear()
        for fp in self.file_paths:
            try:
                doc = Document(fp)
                self._text_cache[fp] = self.get_document_text(doc)
            except Exception:
                self._text_cache[fp] = ""

    def _do_live_count(self):
        """Perform the live count against cached text"""
        self._live_count_after_id = None
        search_for = self.get_processed_search_text()
        if not search_for or not self._text_cache:
            self.match_counter_label.config(text="")
            return
        case_sensitive = self.case_sensitive_var.get()
        use_regex = self.regex_var.get()
        whole_word = self.whole_word_var.get()
        
        if use_regex:
            try:
                re.compile(search_for)
            except re.error:
                self.match_counter_label.config(text="Invalid regex", fg="red")
                return
        
        total = 0
        files_with = 0
        for fp, text in self._text_cache.items():
            c = self.count_occurrences(text, search_for, case_sensitive, use_regex, whole_word)
            total += c
            if c > 0:
                files_with += 1
        
        if total > 0:
            self.match_counter_label.config(
                text=f"Live: {total} match(es) in {files_with} file(s)  [main content only]",
                fg="#1565c0")
        else:
            self.match_counter_label.config(text="Live: No matches found", fg="#888888")

    def check_hyperlinks(self):
        """Check selected files for hyperlinks using python-docx"""
        if not self.file_paths:
            messagebox.showwarning("Warning", "Please add at least one Word document.")
            return
        
        try:
            self.status_label.config(text="Checking for hyperlinks...", fg="orange")
            self.progress_frame.pack(fill=tk.X, pady=(10, 0))
            self.root.update()
            
            hyperlink_results = []
            total_hyperlinks = 0
            files_with_hyperlinks = 0
            
            for i, file_path in enumerate(self.file_paths):
                # Update progress
                progress_percent = int((i / len(self.file_paths)) * 100)
                self.progress_label.config(text=f"Checking {i+1}/{len(self.file_paths)}: {os.path.basename(file_path)}")
                self.progress_percent_label.config(text=f"Progress: {progress_percent}%")
                self.root.update()
                
                try:
                    doc = Document(file_path)
                    file_hyperlinks = []
                    file_hyperlink_count = 0
                    
                    # Check hyperlinks in paragraphs
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            if hasattr(run.element, 'hyperlink') and run.element.hyperlink is not None:
                                hyperlink = run.element.hyperlink
                                # Get hyperlink address
                                if hasattr(hyperlink, 'address') and hyperlink.address:
                                    file_hyperlinks.append({
                                        'text': run.text,
                                        'url': hyperlink.address,
                                        'location': 'paragraph'
                                    })
                                    file_hyperlink_count += 1
                    
                    # Check hyperlinks in tables
                    for table_idx, table in enumerate(doc.tables):
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if hasattr(run.element, 'hyperlink') and run.element.hyperlink is not None:
                                            hyperlink = run.element.hyperlink
                                            if hasattr(hyperlink, 'address') and hyperlink.address:
                                                file_hyperlinks.append({
                                                    'text': run.text,
                                                    'url': hyperlink.address,
                                                    'location': f'table {table_idx+1}, row {row_idx+1}, cell {cell_idx+1}'
                                                })
                                                file_hyperlink_count += 1
                    
                    # Alternative method: Check relationships for hyperlinks
                    if hasattr(doc, 'part') and hasattr(doc.part, 'rels'):
                        for rel in doc.part.rels.values():
                            if rel.reltype.endswith('/hyperlink'):
                                # This is a hyperlink relationship
                                # Try to find the text that uses this relationship
                                file_hyperlink_count += 1
                    
                    hyperlink_results.append({
                        'filename': os.path.basename(file_path),
                        'hyperlink_count': file_hyperlink_count,
                        'hyperlinks': file_hyperlinks[:10],  # Limit to first 10 for display
                        'total_found': len(file_hyperlinks)
                    })
                    
                    if file_hyperlink_count > 0:
                        files_with_hyperlinks += 1
                        total_hyperlinks += file_hyperlink_count
                        
                except Exception as e:
                    hyperlink_results.append({
                        'filename': os.path.basename(file_path),
                        'hyperlink_count': 0,
                        'hyperlinks': [],
                        'error': str(e)
                    })
            
            # Hide progress
            self.progress_frame.pack_forget()
            self.status_label.config(text="Hyperlink check completed!", fg="green")
            
            # Show results
            self.show_hyperlink_results(hyperlink_results, total_hyperlinks, files_with_hyperlinks)
            
        except Exception as e:
            self.progress_frame.pack_forget()
            self.status_label.config(text="Error occurred", fg="red")
            messagebox.showerror("Error", f"Error during hyperlink check: {str(e)}")

    def show_hyperlink_results(self, results, total_hyperlinks, files_with_hyperlinks):
        """Show hyperlink check results in scrollable window"""
        message = "-" * 26 + "\n"
        message += f"🔗 Hyperlink check results\n"
        message += "-" * 26 + "\n\n"
        message += f"📊 Summary:\n"
        message += f"Files analyzed: {len(results)}\n"
        message += f"Total hyperlinks found: {total_hyperlinks}\n"
        message += f"Files with hyperlinks: {files_with_hyperlinks}\n\n"
        message += "=" * 77 + "\n\n"
        
        if total_hyperlinks > 0:
            message += "⚠️  Reminder for hyperlinks:\n"
            message += "• Standard Replace updates display text but removes the actual link!\n"
            message += "• Advanced Replace preserves hyperlinks when editing their display text.\n"
            message += "• If you are not going to replace the display text of a hyperlink, you can still use Standard Replace as it is faster.\n\n"
            message += "=" * 77 + "\n\n"
        
        # Show detailed file results
        for i, result in enumerate(results, 1):
            message += f"📁 File {i}: {result['filename']}\n"
            
            if 'error' in result:
                message += f"   ❌ Error: {result['error']}\n"
            elif result['hyperlink_count'] > 0:
                message += f"   🔗 Hyperlinks found: {result['hyperlink_count']}\n"
                
                # Show sample hyperlinks
                if result['hyperlinks']:
                    message += f"   📋 Sample hyperlinks:\n"
                    for idx, hyperlink in enumerate(result['hyperlinks'][:5], 1):  # Show max 5
                        text_preview = hyperlink['text'][:30] + "..." if len(hyperlink['text']) > 30 else hyperlink['text']
                        url_preview = hyperlink['url'][:40] + "..." if len(hyperlink['url']) > 40 else hyperlink['url']
                        message += f"      {idx}. Text: '{text_preview}'\n"
                        message += f"         URL: {url_preview}\n"
                        message += f"         Location: {hyperlink['location']}\n"
                    
                    if result['total_found'] > 5:
                        message += f"      ... and {result['total_found'] - 5} more hyperlinks\n"
            else:
                message += f"   ✅ No hyperlinks found\n"
            
            message += "\n" + "=" * 77 + "\n\n"
        
        if total_hyperlinks < 1:
            message += "💡 Both Standard and Advanced Replace are safe to use."
        else:
            message += "💡 Only use Standard Replace if you are not going to change display texts of hyperlinks."

        self.show_scrollable_results(message, "Hyperlink check results")

    def handle_delete_key(self, event):
        """Handle Delete key press in file listbox"""
        self.remove_selected_file()
        return "break"  # Prevent any other handling
    
    def show_shortcuts_tooltip(self, event):
        """Show shortcuts tooltip on hover - SMART VISIBILITY MANAGEMENT"""
        # Prevent multiple tooltips - but check if existing one is still valid
        if hasattr(self, 'tooltip'):
            try:
                if self.tooltip.winfo_exists():
                    return
            except tk.TclError:
                # Tooltip was destroyed but attribute still exists
                del self.tooltip

        # Create tooltip window - COMPLETELY INDEPENDENT
        self.tooltip = tk.Toplevel()  # Don't pass self.root as parent!
        self.tooltip.wm_overrideredirect(True)  # Remove window decorations
        self.tooltip.configure(bg="lightyellow", relief="solid", borderwidth=1)
        
        # Make tooltip stay on top but don't make it transient
        self.tooltip.attributes('-topmost', True)
        
        # BIND FOCUS EVENTS TO MAIN WINDOW TO CONTROL TOOLTIP VISIBILITY
        def on_main_focus_in(event):
            """Show tooltip when main window gets focus"""
            if hasattr(self, 'tooltip'):
                try:
                    self.tooltip.deiconify()  # Show tooltip
                except tk.TclError:
                    pass
        
        def on_main_focus_out(event):
            """Hide tooltip when main window loses focus"""
            if hasattr(self, 'tooltip'):
                try:
                    self.tooltip.withdraw()  # Hide tooltip (but don't destroy)
                except tk.TclError:
                    pass
        
        # Bind focus events to main window
        self.root.bind('<FocusIn>', on_main_focus_in, add='+')
        self.root.bind('<FocusOut>', on_main_focus_out, add='+')
        
        # Also bind to window state changes (minimize/restore)
        def on_window_state_change(event):
            """Handle window minimize/restore"""
            if hasattr(self, 'tooltip'):
                try:
                    if self.root.state() == 'iconic':  # Minimized
                        self.tooltip.withdraw()
                    else:  # Normal or zoomed
                        self.tooltip.deiconify()
                except tk.TclError:
                    pass
        
        self.root.bind('<Unmap>', on_window_state_change, add='+')
        self.root.bind('<Map>', on_window_state_change, add='+')
        
        # Create a frame to hold everything
        main_frame = tk.Frame(self.tooltip, bg="lightyellow")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Tooltip content
        shortcuts_text = """    𝗥𝗲𝗽𝗹𝗮𝗰𝗲𝗺𝗲𝗻𝘁 𝗺𝗲𝘁𝗵𝗼𝗱𝘀 
    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    Preview:
    • Shows what will be changed

    Standard Replace:
    • Handles standard text & tables (fast)
    • ⚠️ Hyperlinks: Editing the display text of a hyperlink updates the display text, but removes the link.

    Advanced Replace:
    • Handles all document areas (complete coverage) and handles hyperlinks correctly

    𝗞𝗲𝘆𝗯𝗼𝗮𝗿𝗱 𝘀𝗵𝗼𝗿𝘁𝗰𝘂𝘁𝘀
    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    Actions:
    • F5 - Preview
    • Ctrl+Enter - Standard Replace
    • Shift+Enter - Advanced Replace
    
    Files:
    • Ctrl+O - Add more files
    • Delete - Remove selected files (when file list is focused)
    
    Navigation:
    • Ctrl+1 - Focus Search box
    • Ctrl+2 - Focus Replace box
    • Tab - Switch between boxes

    Text Editing:
    • Ctrl+Tab - Add a tab (\\t) in search/replace box
    • Ctrl+Z - Undo
    • Ctrl+Y - Redo
    • Shift+Space - Insert _nbsp_ at cursor position

    Exit:
    • Escape/Ctrl+Q - Close application

    𝗠𝗶𝘀𝗰𝗲𝗹𝗹𝗮𝗻𝗲𝗼𝘂𝘀
    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    Non-breaking spaces:
    • Type '_nbsp_' for non-breaking spaces or use Shift+Space
    • 🔍 nbsp check - Counts non-breaking spaces in search/replace boxes

    Hyperlinks:
    • 🔗 Hyperlink check - Checks the selected files for hyperlinks

    Regex:
    • Supported in Standard Replace/Preview only (Python regex syntax)
    • Not supported in Advanced Replace/Preview (uses Word's native engine)
    
    𝗙𝗼𝗿𝘇𝗮 𝗜𝗻𝘁𝗲𝗿!"""
        
        # Create scrollable text widget instead of label
        text_widget = tk.Text(main_frame, 
                            font=("Arial", 9), 
                            bg="lightyellow", 
                            wrap=tk.WORD,
                            width=80,  # Set reasonable width
                            height=30,  # Set reasonable height
                            relief="flat",
                            borderwidth=0,
                            padx=5,
                            pady=3)
        
        # Create scrollbar
        scrollbar = tk.Scrollbar(main_frame, orient="vertical", command=text_widget.yview)
        text_widget.config(yscrollcommand=scrollbar.set)
        
        # Pack text widget and scrollbar
        text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Insert text and make read-only
        text_widget.insert("1.0", shortcuts_text)
        text_widget.config(state=tk.DISABLED)
        
        # SIMPLE BUT EFFECTIVE POSITIONING - STAY NEAR MAIN WINDOW
        # Get main window bounds
        main_x = self.root.winfo_rootx()
        main_y = self.root.winfo_rooty()
        main_width = self.root.winfo_width()
        main_height = self.root.winfo_height()
        
        # Get icon position
        icon_x = self.info_icon.winfo_rootx()
        icon_y = self.info_icon.winfo_rooty()
        
        # Set a reasonable tooltip size first
        self.tooltip.geometry("650x600")  # Fixed size
        self.tooltip.update_idletasks()
        
        tooltip_width = 650
        tooltip_height = 600
        
        # Strategy: Try positions in order of preference, all relative to main window
        positions_to_try = [
            # 1. Left of icon (preferred)
            (icon_x - tooltip_width - 10, icon_y),
            # 2. Right of main window
            (main_x + main_width + 10, main_y),
            # 3. Left of main window
            (main_x - tooltip_width - 10, main_y),
            # 4. Above main window
            (main_x, main_y - tooltip_height - 10),
            # 5. Below main window
            (main_x, main_y + main_height + 10),
            # 6. Centered on main window (last resort)
            (main_x + (main_width - tooltip_width) // 2, main_y + (main_height - tooltip_height) // 2)
        ]
        
        # Find the first position that keeps tooltip reasonably close to main window
        final_x, final_y = positions_to_try[0]  # Default to first option
        
        for test_x, test_y in positions_to_try:
            # Check if this position keeps tooltip close to main window
            distance_from_main = abs(test_x - main_x) + abs(test_y - main_y)
            if distance_from_main < 1000:  # Within reasonable distance
                final_x, final_y = test_x, test_y
                break
        
        # Ensure tooltip doesn't go off-screen (but prioritize staying near main window)
        # Only adjust if absolutely necessary
        if final_x < main_x - 800:  # Too far left
            final_x = main_x - 400
        if final_y < main_y - 400:  # Too far up
            final_y = main_y - 200
        
        # Apply final position
        self.tooltip.geometry(f"650x600+{final_x}+{final_y}")
        
        # Add close button in the top-right corner - ONLY WAY TO CLOSE
        close_btn = tk.Button(main_frame, text="🗙", 
                            command=self.hide_shortcuts_tooltip_manual,
                            font=("Arial", 8, "bold"),
                            bg="lightcoral", fg="white",
                            width=2, height=1,
                            relief="flat")
        close_btn.place(relx=1.0, rely=0.0, anchor="ne", x=-23, y=0)

    def hide_shortcuts_tooltip(self, event):
        """Do nothing - tooltip only closes when X is clicked"""
        pass

    def hide_shortcuts_tooltip_manual(self):
        """Manually hide tooltip when close button is clicked"""
        if hasattr(self, 'tooltip'):
            try:
                self.tooltip.destroy()
            except tk.TclError:
                pass  # Already destroyed
            del self.tooltip



#########END OF PART 1######




#########PART 2######
    def setup_keyboard_bindings(self):
        """Setup custom keyboard bindings"""
        # Existing text box bindings
        self.search_text.bind('<Tab>', self.focus_replace_text)
        self.replace_text.bind('<Tab>', self.focus_search_text)
        self.search_text.bind('<Control-Tab>', self.insert_tab)
        self.replace_text.bind('<Control-Tab>', self.insert_tab)
        self.search_text.bind('<Control-z>', self.undo_search)
        self.search_text.bind('<Control-y>', self.redo_search)
        self.replace_text.bind('<Control-z>', self.undo_replace)
        self.replace_text.bind('<Control-y>', self.redo_replace)
        self.search_text.bind('<Control-Shift-Z>', self.redo_search)
        self.replace_text.bind('<Control-Shift-Z>', self.redo_replace)

        # NEW: NBSP shortcut - Shift+Space inserts _nbsp_ at cursor position
        self.search_text.bind('<Shift-space>', self.insert_nbsp_placeholder)
        self.replace_text.bind('<Shift-space>', self.insert_nbsp_placeholder)
        
        # GLOBAL KEYBOARD SHORTCUTS (bind to root)
        self.root.bind('<Control-Return>', lambda e: self.replace_text_in_documents())
        self.root.bind('<F5>', lambda e: self.preview_changes())
        self.root.bind('<Control-o>', lambda e: self.add_files())
        self.root.bind('<Control-O>', lambda e: self.add_files())
        self.root.bind('<Escape>', lambda e: self.root.destroy())
        self.root.bind('<Control-q>', lambda e: self.root.destroy())
        self.root.bind('<Control-Q>', lambda e: self.root.destroy())
        self.root.bind('<Control-Shift-Return>', lambda e: self.replace_and_close())
        self.root.bind('<Control-Key-1>', lambda e: self.search_text.focus())
        self.root.bind('<Control-Key-2>', lambda e: self.replace_text.focus())
        self.root.bind('<Shift-Return>', lambda e: self.advanced_replace_with_vba())
        
        # BIND SHORTCUTS TO TEXT WIDGETS TOO (to prevent line breaks)
        # Search text shortcuts
        self.search_text.bind('<Control-Return>', self.handle_shortcut_replace)
        self.search_text.bind('<Control-Shift-Return>', self.handle_shortcut_replace_close)
        self.search_text.bind('<F5>', self.handle_shortcut_preview)
        self.search_text.bind('<Control-o>', self.handle_shortcut_add_files)
        self.search_text.bind('<Control-O>', self.handle_shortcut_add_files)
        self.search_text.bind('<Escape>', self.handle_shortcut_close)
        self.search_text.bind('<Control-q>', self.handle_shortcut_close)
        self.search_text.bind('<Control-Q>', self.handle_shortcut_close)
        self.search_text.bind('<Shift-Return>', self.handle_shortcut_advanced_replace)
        
        # Replace text shortcuts
        self.replace_text.bind('<Control-Return>', self.handle_shortcut_replace)
        self.replace_text.bind('<Control-Shift-Return>', self.handle_shortcut_replace_close)
        self.replace_text.bind('<F5>', self.handle_shortcut_preview)
        self.replace_text.bind('<Control-o>', self.handle_shortcut_add_files)
        self.replace_text.bind('<Control-O>', self.handle_shortcut_add_files)
        self.replace_text.bind('<Escape>', self.handle_shortcut_close)
        self.replace_text.bind('<Control-q>', self.handle_shortcut_close)
        self.replace_text.bind('<Control-Q>', self.handle_shortcut_close)
        self.replace_text.bind('<Shift-Return>', self.handle_shortcut_advanced_replace)
        
        # Make sure the window can receive focus for global shortcuts
        self.root.focus_set()
    
    def handle_shortcut_replace(self, event):
        """Handle Ctrl+Enter shortcut from text widgets"""
        self.replace_text_in_documents()
        return "break"  # Prevent line break insertion
    
    def handle_shortcut_replace_close(self, event):
        """Handle Ctrl+Shift+Enter shortcut from text widgets"""
        self.replace_and_close()
        return "break"  # Prevent line break insertion
    
    def handle_shortcut_preview(self, event):
        """Handle F5 shortcut from text widgets"""
        self.preview_changes()
        return "break"  # Prevent any default behavior
    
    def handle_shortcut_add_files(self, event):
        """Handle Ctrl+O shortcut from text widgets"""
        self.add_files()
        return "break"  # Prevent any default behavior
    
    def handle_shortcut_close(self, event):
        """Handle Escape/Ctrl+Q shortcut from text widgets"""
        self.root.destroy()
        return "break"  # Prevent any default behavior

    def handle_shortcut_advanced_replace(self, event):
        """Handle Shift+Enter shortcut from text widgets"""
        self.advanced_replace_with_vba()
        return "break"  # Prevent line break insertion
    
    def replace_and_close(self):
        """Replace and close application"""
        if self.replace_text_in_documents(close_after=True):
            pass  # The close will be handled in the results window
    
    def insert_tab(self, event):
        """Insert a tab character at cursor position"""
        widget = event.widget
        widget.insert(tk.INSERT, '\t')
        return "break"  # Prevent any other handling

    def insert_nbsp_placeholder(self, event):
        """Insert _nbsp_ placeholder at cursor position"""
        widget = event.widget
        cursor_pos = widget.index(tk.INSERT)
        widget.insert(cursor_pos, '_nbsp_')
        return "break"  # Prevent any other handling
    
    def focus_replace_text(self, event):
        """Move focus from search to replace text box"""
        self.replace_text.focus()
        return "break"  # Prevent default Tab behavior
    
    def focus_search_text(self, event):
        """Move focus from replace to search text box"""
        self.search_text.focus()
        return "break"  # Prevent default Tab behavior
    
    def undo_search(self, event):
        """Undo in search text box"""
        try:
            self.search_text.edit_undo()
        except tk.TclError:
            pass  # No more undo operations
        return "break"
    
    def redo_search(self, event):
        """Redo in search text box"""
        try:
            self.search_text.edit_redo()
        except tk.TclError:
            pass  # No more redo operations
        return "break"
    
    def undo_replace(self, event):
        """Undo in replace text box"""
        try:
            self.replace_text.edit_undo()
        except tk.TclError:
            pass  # No more undo operations
        return "break"
    
    def redo_replace(self, event):
        """Redo in replace text box"""
        try:
            self.replace_text.edit_redo()
        except tk.TclError:
            pass  # No more redo operations
        return "break"
    
    def paste_to_search(self):
        """Paste clipboard content to search text box"""
        try:
            clipboard_content = self.root.clipboard_get()
            self.search_text.delete("1.0", tk.END)
            self.search_text.insert("1.0", clipboard_content)
            self.status_label.config(text="Pasted clipboard content to search box.", fg="blue")
            self._schedule_live_count()
        except tk.TclError:
            self.status_label.config(text="Clipboard is empty or contains non-text data.", fg="orange")
    
    def paste_to_replace(self):
        """Paste clipboard content to replace text box"""
        try:
            clipboard_content = self.root.clipboard_get()
            self.replace_text.delete("1.0", tk.END)
            self.replace_text.insert("1.0", clipboard_content)
            self.status_label.config(text="Pasted clipboard content to replace box.", fg="blue")
        except tk.TclError:
            self.status_label.config(text="Clipboard is empty or contains non-text data.", fg="orange")
    
    def update_title(self):
        """Update window title with file count"""
        file_count = len(self.file_paths)
        if file_count == 0:
            self.root.title("Bulk Text Replacement for Word - No files selected.")
        else:
            self.root.title(f"Bulk Text Replacement for Word - {file_count} file{'s' if file_count != 1 else ''}")
    
    def update_file_list(self):
        """Update the file listbox and related UI elements"""
        self.file_listbox.delete(0, tk.END)
        
        for file_path in self.file_paths:
            self.file_listbox.insert(tk.END, os.path.basename(file_path))
        
        # Update file count label
        file_count = len(self.file_paths)
        if file_count == 0:
            self.file_count_label.config(text="No files selected.")
            self.replace_btn.config(state="disabled")
        else:
            self.file_count_label.config(text=f"{file_count} file{'s' if file_count != 1 else ''} selected.")
            self.replace_btn.config(state="normal")
        
        self.update_title()
    
    def add_files(self):
        """Add more Word documents to the list"""
        file_types = [
            ("Word Documents", "*.docx *.doc *.docm"),
            ("Word 2007+ Documents", "*.docx *.docm"),
            ("Word 97-2003 Documents", "*.doc"),
            ("All files", "*.*")
        ]
        
        selected_files = filedialog.askopenfilenames(
            title="Select Word Documents to add",
            filetypes=file_types,
            initialdir=os.path.dirname(self.file_paths[0]) if self.file_paths else None
        )
        
        added_count = 0
        for file_path in selected_files:
            # Normalize the path to avoid duplicates due to different path formats
            normalized_path = os.path.normpath(os.path.abspath(file_path))
            
            # Check if this normalized path is already in our list (also normalized)
            already_exists = False
            for existing_path in self.file_paths:
                if os.path.normpath(os.path.abspath(existing_path)) == normalized_path:
                    already_exists = True
                    break
            
            if not already_exists:
                self.file_paths.append(file_path)
                added_count += 1
        
        if added_count > 0:
            self._refresh_text_cache()
            self.update_file_list()
            self.status_label.config(text=f"Added {added_count} file{'s' if added_count != 1 else ''}", fg="blue")
            self._schedule_live_count()
        else:
            if len(selected_files) > 0:
                self.status_label.config(text="Files already in list - no duplicates added.", fg="orange")
            else:
                self.status_label.config(text="No new files added.", fg="orange")
    
    def remove_selected_file(self):
        """Remove the selected file(s) from the list"""
        selections = self.file_listbox.curselection()
        if selections:
            # Get filenames before removing (for status message)
            removed_files = []
            for index in selections:
                removed_files.append(os.path.basename(self.file_paths[index]))
            
            # Remove files in reverse order to maintain correct indices
            for index in reversed(selections):
                del self.file_paths[index]
            
            self._refresh_text_cache()
            self.update_file_list()
            self._schedule_live_count()
            
            # Update status message
            if len(removed_files) == 1:
                self.status_label.config(text=f"Removed: {removed_files[0]}", fg="orange")
            else:
                self.status_label.config(text=f"Removed {len(removed_files)} files", fg="orange")
        else:
            messagebox.showwarning("Warning", "Please select one or more files to remove.")
    
    def clear_all_files(self):
        """Clear all files from the list"""
        if self.file_paths:
            if messagebox.askyesno("Confirm", "Remove all files from the list?"):
                file_count = len(self.file_paths)
                self.file_paths.clear()
                self._text_cache.clear()
                self.update_file_list()
                self.match_counter_label.config(text="")
                self.status_label.config(text=f"Cleared {file_count} file{'s' if file_count != 1 else ''}", fg="orange")
    
    def get_document_text(self, doc):
        """Extract all text from document including paragraphs and tables"""
        full_text = []
        
        # Get text from paragraphs (strip invisible formatting chars)
        for paragraph in doc.paragraphs:
            full_text.append(self._strip_invisible_chars(paragraph.text))
        
        # Get text from tables (including nested tables)
        for table in doc.tables:
            self._collect_table_text(table, full_text)
        
        return '\n'.join(full_text)
    
    def _collect_table_text(self, table, text_list):
        """Recursively collect text from a table and its nested tables"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text_list.append(self._strip_invisible_chars(paragraph.text))
                for nested_table in cell.tables:
                    self._collect_table_text(nested_table, text_list)
    
    def count_occurrences(self, text, search_for, case_sensitive=False, use_regex=False, whole_word=False):
        """Count occurrences with case sensitivity, regex, and whole word options"""
        if not search_for:
            return 0
        # Normalize: strip invisible formatting characters
        text = self._strip_invisible_chars(text)
        search_for = self._strip_invisible_chars(search_for)
        if use_regex:
            try:
                flags = 0 if case_sensitive else re.IGNORECASE
                return len(re.findall(search_for, text, flags))
            except re.error:
                return 0
        if whole_word:
            pattern = r'\b' + re.escape(search_for) + r'\b'
            flags = 0 if case_sensitive else re.IGNORECASE
            return len(re.findall(pattern, text, flags))
        if not case_sensitive:
            return text.lower().count(search_for.lower())
        else:
            return text.count(search_for)

    def _find_match_contexts(self, text, search_for, case_sensitive=False, use_regex=False, whole_word=False, context_chars=40, max_matches=5):
        """Extract surrounding context for each match (for diff preview)"""
        contexts = []
        if not search_for:
            return contexts
        text = self._strip_invisible_chars(text)
        search_for = self._strip_invisible_chars(search_for)
        if use_regex:
            pattern = search_for
        elif whole_word:
            pattern = r'\b' + re.escape(search_for) + r'\b'
        else:
            pattern = re.escape(search_for)
        flags = 0 if case_sensitive else re.IGNORECASE
        try:
            for i, m in enumerate(re.finditer(pattern, text, flags)):
                if i >= max_matches:
                    break
                start = max(0, m.start() - context_chars)
                end = min(len(text), m.end() + context_chars)
                before = text[start:m.start()].replace('\n', ' ').replace('\r', '')
                matched = m.group().replace('\n', ' ').replace('\r', '')
                after = text[m.end():end].replace('\n', ' ').replace('\r', '')
                prefix = "..." if start > 0 else ""
                suffix = "..." if end < len(text) else ""
                contexts.append(f'{prefix}{before}[{matched}]{after}{suffix}')
        except re.error:
            pass
        return contexts

    def preview_changes(self):
        """Show preview with choice between Standard (fast) or Advanced (comprehensive) mode"""
        if not self.file_paths:
            messagebox.showwarning("Warning", "Please add at least one Word document.")
            return
            
        search_for = self.get_processed_search_text()
        
        if not search_for:
            messagebox.showwarning("Warning", "Please enter text to search for.")
            return
        
        use_regex = self.regex_var.get()
        if use_regex:
            try:
                re.compile(search_for)
            except re.error as e:
                messagebox.showerror("Invalid Regex", f"Invalid regex pattern:\n{str(e)}")
                return
        
        # NEW: Ask user for preview mode
        preview_choice = messagebox.askyesnocancel(
            "Preview Mode Selection",
            "Do you want to run the fast Standard Preview?\n\n" +
            "YES = Standard Preview (Fast)\n" +
            "   • Main content only (paragraphs + tables)\n\n" +
            "NO = Advanced Preview (Comprehensive)\n" +
            "   • Main content + shapes + headers + footers + footnotes + endnotes + form fields + hyperlinks"
        )
        
        if preview_choice is None:  # User clicked Cancel
            return
        elif preview_choice:  # User clicked Yes - Standard Preview
            self.preview_standard_only()
        else:  # User clicked No - Advanced Preview
            self.preview_comprehensive()

    def preview_standard_only(self):
        """Fast preview - Main content only (python-docx)"""
        search_for = self.get_processed_search_text()
        case_sensitive = self.case_sensitive_var.get()
        use_regex = self.regex_var.get()
        whole_word = self.whole_word_var.get()
        
        try:
            preview_details = []
            total_matches = 0
            files_with_matches = 0
            
            # Show progress for standard preview
            self.status_label.config(text="Standard preview in progress...", fg="orange")
            self.progress_frame.pack(fill=tk.X, pady=(10, 0))
            self.root.update()
            
            for i, file_path in enumerate(self.file_paths):
                # Update progress
                progress_percent = int((i / len(self.file_paths)) * 100)
                self.progress_label.config(text=f"Analyzing {i+1}/{len(self.file_paths)}: {os.path.basename(file_path)}")
                self.progress_percent_label.config(text=f"Progress: {progress_percent}%")
                self.root.update()
                
                file_result = {
                    'filename': os.path.basename(file_path),
                    'main_total': 0,
                    'details': []
                }
                
                try:
                    # Analyze main content only (python-docx)
                    doc = Document(file_path)
                    document_text = self.get_document_text(doc)
                    main_occurrences = self.count_occurrences(document_text, search_for, case_sensitive, use_regex, whole_word)
                    
                    file_result['main_total'] = main_occurrences
                    file_result['contexts'] = []
                    if main_occurrences > 0:
                        files_with_matches += 1
                        total_matches += main_occurrences
                        file_result['details'].append(f"  📄 Main content: {main_occurrences} match(es)")
                        file_result['contexts'] = self._find_match_contexts(
                            document_text, search_for, case_sensitive, use_regex, whole_word)
                    else:
                        file_result['details'].append(f"  📄 Main content: No matches")
                    
                    preview_details.append(file_result)
                    
                except Exception as e:
                    file_result['details'] = [f"  ❌ Error: {str(e)}"]
                    preview_details.append(file_result)
            
            # Hide progress
            self.progress_frame.pack_forget()
            self.status_label.config(text="Standard preview completed!", fg="green")
            
            # Show standard preview results
            self.show_standard_preview_results(
                preview_details, search_for, self.get_processed_replace_text(),
                total_matches, files_with_matches, case_sensitive
            )
            
        except Exception as e:
            self.progress_frame.pack_forget()
            self.status_label.config(text="Error occurred", fg="red")
            messagebox.showerror("Error", f"Error during standard preview: {str(e)}")

    def show_standard_preview_results(self, results, search_text, replace_text, total_matches, files_with_matches, case_sensitive):
        """Show standard preview results (main content only)"""
        case_info = " (case sensitive)" if case_sensitive else " (case insensitive)"
        
        message = "-" * 46 + "\n"
        message += f"📄 Standard Preview results{case_info}\n"
        message += "-" * 46 + "\n\n"
        message += f"⚡ Fast Analysis: {total_matches} match(es) found in main content\n\n"
        message += f"📊 Summary:\n"
        message += f"Files analyzed: {len(results)}\n"
        message += f"Main content matches: {total_matches} (in {files_with_matches} files)\n"
        message += f"Search text: '{search_text[:40]}{'...' if len(search_text) > 40 else ''}'\n"
        message += f"Replace text: '{replace_text[:40]}{'...' if len(replace_text) > 40 else ''}'\n"
        message += f"Coverage: Main text and tables only\n"
        message += "=" * 77 + "\n\n"
        
        # Show detailed file results
        for i, result in enumerate(results, 1):
            message += f"📁 FILE {i}: {result['filename']}\n"
            
            if result['main_total'] > 0:
                message += f"   ✅ Matches found: {result['main_total']}\n"
            else:
                message += f"   ✅ No matches found\n"
            
            for detail in result['details']:
                message += f"   {detail}\n"
            
            # Diff preview: show context snippets
            contexts = result.get('contexts', [])
            if contexts:
                message += f"\n   📝 Preview (first {len(contexts)} match{'es' if len(contexts) != 1 else ''}):\n"
                for j, ctx in enumerate(contexts, 1):
                    message += f"      {j}. {ctx}\n"
                if result['main_total'] > len(contexts):
                    message += f"      ... and {result['main_total'] - len(contexts)} more\n"
            
            message += "\n" + "=" * 77 + "\n\n"
        
        # Recommendations
        if total_matches > 0:
            message += "💡 Replacement recommendations:\n"
            message += "Use 'Standard Replace' for fast replacement of main content\n"
            message += "Use 'Advanced Replace' if you also need shapes/headers/footers\n"
        else:
            message += "💡 Tip: No matches found in main content. Try 'Advanced Preview' to run a comprehensive check."
        
        self.show_scrollable_results(message, "Standard Preview results")

    def preview_comprehensive(self):
        """Comprehensive preview - All areas using Word COM automation"""
        search_for = self.get_processed_search_text()
        case_sensitive = self.case_sensitive_var.get()
        
        if not HAS_WIN32COM:
            messagebox.showerror("Error", 
                "Advanced preview requires the 'pywin32' package.\n\n"
                "Install it with: pip install pywin32")
            return
        
        if self.regex_var.get():
            messagebox.showinfo("Regex Not Supported", 
                "Regex is only supported with Standard Preview/Replace.\n\n"
                "Advanced Preview uses Word's native engine\n"
                "which does not support Python regex syntax.\n\n"
                "Please uncheck 'Regex' or use Standard Preview instead.")
            return
        
        word_app = None
        try:
            preview_details = []
            total_advanced_matches = 0
            files_with_advanced_matches = 0
            
            # Show progress for comprehensive preview
            self.status_label.config(text="Comprehensive analysis in progress...", fg="orange")
            self.progress_frame.pack(fill=tk.X, pady=(10, 0))
            self.root.update()
            
            # Open Word once for all files
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False
            word_app.ScreenUpdating = False
            
            for i, file_path in enumerate(self.file_paths):
                # Update progress
                progress_percent = int((i / len(self.file_paths)) * 100)
                self.progress_label.config(text=f"Analyzing {i+1}/{len(self.file_paths)}: {os.path.basename(file_path)}")
                self.progress_percent_label.config(text=f"Progress: {progress_percent}%")
                self.root.update()
                
                try:
                    file_result = {
                        'filename': os.path.basename(file_path),
                        'total': 0,
                        'details': []
                    }
                    
                    advanced_result = self.preview_advanced_areas(file_path, search_for, case_sensitive, word_app)
                    file_result['total'] = advanced_result['total']
                    
                    if advanced_result['total'] > 0:
                        files_with_advanced_matches += 1
                        total_advanced_matches += advanced_result['total']
                        file_result['details'].extend(advanced_result['details'])
                    else:
                        file_result['details'].append(f"  📄 No matches found in any areas")
                    
                    preview_details.append(file_result)
                    
                except Exception as e:
                    file_result = {
                        'filename': os.path.basename(file_path),
                        'total': 0,
                        'details': [f"  ❌ Error: {str(e)}"]
                    }
                    preview_details.append(file_result)
            
            # Hide progress
            self.progress_frame.pack_forget()
            self.status_label.config(text="Comprehensive preview completed!", fg="green")
            
            # Show comprehensive preview results
            self.show_comprehensive_preview_results(
                preview_details, search_for, self.get_processed_replace_text(),
                total_advanced_matches, files_with_advanced_matches, 
                case_sensitive
            )
            
        except Exception as e:
            self.progress_frame.pack_forget()
            self.status_label.config(text="Error occurred", fg="red")
            messagebox.showerror("Error", f"Error during comprehensive preview: {str(e)}")
        finally:
            try:
                if word_app:
                    word_app.ScreenUpdating = True
                    word_app.Quit()
            except Exception:
                pass
#########END OF PART 2######





#########PART 3A######
    def _build_shape_ranges(self, doc):
        """Pre-build list of (start, end) for all shapes with text frames.
        Used to check if a hyperlink is inside a shape in O(n) instead of O(n*m)."""
        ranges = []
        self._collect_shape_ranges(doc.Shapes, ranges)
        return ranges

    def _collect_shape_ranges(self, shapes, ranges):
        """Recursively collect text frame ranges from shapes (including grouped shapes)"""
        for shape in shapes:
            try:
                if shape.HasTextFrame and shape.TextFrame.HasText:
                    ranges.append((shape.TextFrame.TextRange.Start, shape.TextFrame.TextRange.End))
            except Exception:
                pass
            try:
                if shape.Type == 6:  # msoGroup
                    self._collect_shape_ranges(shape.GroupItems, ranges)
            except Exception:
                pass

    def _is_in_shape_ranges(self, start, end, shape_ranges):
        """Check if a range is inside any pre-cached shape range."""
        for s_start, s_end in shape_ranges:
            if start >= s_start and end <= s_end:
                return True
        return False

    def _count_in_shapes(self, shapes, search_text, case_sensitive, whole_word):
        """Recursively count occurrences in shapes (including grouped shapes)"""
        count = 0
        for shape in shapes:
            try:
                if shape.HasTextFrame and shape.TextFrame.HasText:
                    shape_text = shape.TextFrame.TextRange.Text
                    if len(shape_text) > 1:
                        count += self.count_occurrences(shape_text, search_text, case_sensitive, False, whole_word)
            except Exception:
                pass
            try:
                if shape.Type == 6:  # msoGroup
                    count += self._count_in_shapes(shape.GroupItems, search_text, case_sensitive, whole_word)
            except Exception:
                pass
        return count

    def _replace_in_shapes(self, shapes, search_text, replace_text, case_sensitive, whole_word):
        """Recursively replace in shapes using iterative Find (including grouped shapes)"""
        count = 0
        for shape in shapes:
            try:
                if shape.HasTextFrame and shape.TextFrame.HasText:
                    if len(shape.TextFrame.TextRange.Text) > 1:
                        count += self._find_replace_count(
                            shape.TextFrame.TextRange, search_text, replace_text, case_sensitive, whole_word)
            except Exception:
                pass
            try:
                if shape.Type == 6:  # msoGroup
                    count += self._replace_in_shapes(shape.GroupItems, search_text, replace_text, case_sensitive, whole_word)
            except Exception:
                pass
        return count

    def _find_replace_count(self, rng, search_text, replace_text, case_sensitive, whole_word=False):
        """Iterative Find.Execute (Option A) — replaces one match at a time, returns exact count."""
        count = 0
        rng.Find.ClearFormatting()
        rng.Find.Replacement.ClearFormatting()
        while True:
            found = rng.Find.Execute(
                FindText=search_text, ReplaceWith=replace_text,
                Replace=1, Forward=True, Wrap=0,
                MatchCase=case_sensitive, MatchWholeWord=whole_word,
                MatchWildcards=False, MatchSoundsLike=False,
                MatchAllWordForms=False, Format=False
            )
            if not found:
                break
            count += 1
        return count

    def preview_advanced_areas(self, file_path, search_text, case_sensitive=False, word_app=None):
        """Preview advanced areas using Word COM automation - READ ONLY"""
        result = {
            'total': 0,
            'details': []
        }
        
        whole_word = self.whole_word_var.get()
        own_word_app = False
        doc = None
        try:
            if word_app is None:
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = False
                word_app.ScreenUpdating = False
                own_word_app = True
            
            full_path = os.path.abspath(file_path)
            doc = word_app.Documents.Open(full_path, ReadOnly=True)
            
            shapes_count = 0
            headers_count = 0
            footers_count = 0
            footnotes_count = 0
            endnotes_count = 0
            form_fields_count = 0
            hyperlinks_count = 0
            main_content_count = 0
            
            # Count in text boxes / shapes via StoryRanges (wdTextFrameStory = 5)
            # This catches ALL text boxes regardless of how they were inserted
            try:
                story = doc.StoryRanges(5)  # wdTextFrameStory
                while story:
                    story_text = story.Text
                    if len(story_text) > 1:
                        shapes_count += self.count_occurrences(story_text, search_text, case_sensitive, False, whole_word)
                    try:
                        story = story.NextStoryRange
                    except Exception:
                        break
            except Exception:
                pass
            
            # Count in headers and footers
            for section in doc.Sections:
                for i in range(1, 4):
                    try:
                        if section.Headers(i).Exists:
                            header_text = section.Headers(i).Range.Text
                            if len(header_text) > 1:
                                headers_count += self.count_occurrences(header_text, search_text, case_sensitive, False, whole_word)
                    except Exception:
                        pass
                    try:
                        if section.Footers(i).Exists:
                            footer_text = section.Footers(i).Range.Text
                            if len(footer_text) > 1:
                                footers_count += self.count_occurrences(footer_text, search_text, case_sensitive, False, whole_word)
                    except Exception:
                        pass
            
            # Count in footnotes
            for footnote in doc.Footnotes:
                try:
                    footnote_text = footnote.Range.Text
                    if len(footnote_text) > 1:
                        footnotes_count += self.count_occurrences(footnote_text, search_text, case_sensitive, False, whole_word)
                except Exception:
                    pass
            
            # Count in endnotes
            for endnote in doc.Endnotes:
                try:
                    endnote_text = endnote.Range.Text
                    if len(endnote_text) > 1:
                        endnotes_count += self.count_occurrences(endnote_text, search_text, case_sensitive, False, whole_word)
                except Exception:
                    pass
            
            # Count in form fields
            for field in doc.FormFields:
                try:
                    if field.Type == 70:  # wdFieldFormTextInput
                        field_text = field.Result
                        if len(field_text) > 0:
                            form_fields_count += self.count_occurrences(field_text, search_text, case_sensitive, False, whole_word)
                except Exception:
                    pass
            
            # Count in hyperlinks (ONLY those NOT in shapes to avoid double counting)
            shape_ranges = self._build_shape_ranges(doc)
            for hyperlink in doc.Hyperlinks:
                try:
                    display_text = hyperlink.TextToDisplay
                    if len(display_text) > 0:
                        if not self._is_in_shape_ranges(hyperlink.Range.Start, hyperlink.Range.End, shape_ranges):
                            hyperlinks_count += self.count_occurrences(display_text, search_text, case_sensitive, False, whole_word)
                except Exception:
                    pass
            
            # Count in main content (subtract hyperlink matches to avoid double-counting,
            # since doc.Content.Text includes hyperlink display text)
            try:
                main_text = doc.Content.Text
                if len(main_text) > 1:
                    raw_main_count = self.count_occurrences(main_text, search_text, case_sensitive, False, whole_word)
                    main_content_count = max(0, raw_main_count - hyperlinks_count)
            except Exception:
                pass
            
            # Build results
            if shapes_count > 0:
                result['details'].append(f"  📦 Text boxes: {shapes_count} match(es)")
                result['total'] += shapes_count
            if headers_count > 0:
                result['details'].append(f"  📄 Headers: {headers_count} match(es)")
                result['total'] += headers_count
            if footers_count > 0:
                result['details'].append(f"  📄 Footers: {footers_count} match(es)")
                result['total'] += footers_count
            if footnotes_count > 0:
                result['details'].append(f"  📝 Footnotes: {footnotes_count} match(es)")
                result['total'] += footnotes_count
            if endnotes_count > 0:
                result['details'].append(f"  📝 Endnotes: {endnotes_count} match(es)")
                result['total'] += endnotes_count
            if form_fields_count > 0:
                result['details'].append(f"  📋 Form fields: {form_fields_count} match(es)")
                result['total'] += form_fields_count
            if hyperlinks_count > 0:
                result['details'].append(f"  🔗 Hyperlinks: {hyperlinks_count} match(es)")
                result['total'] += hyperlinks_count
            if main_content_count > 0:
                result['details'].append(f"  📄 Main content: {main_content_count} match(es)")
                result['total'] += main_content_count
            
            return result
            
        except Exception as e:
            result['details'] = [f"  ❌ Advanced preview error: {str(e)}"]
            return result
        finally:
            try:
                if doc:
                    doc.Close(False)
            except Exception:
                pass
            if own_word_app:
                try:
                    word_app.ScreenUpdating = True
                    word_app.Quit()
                except Exception:
                    pass

    def show_comprehensive_preview_results(self, results, search_text, replace_text, 
                                         total_matches, files_with_matches, 
                                         case_sensitive):
        """Show comprehensive preview results"""
        case_info = " (case sensitive)" if case_sensitive else " (case insensitive)"

        message = "-" * 46 + "\n"    
        message += f"🔍 Advanced Preview results{case_info}\n"
        message += "-" * 46 + "\n\n" 
        message += f"🎯 Complete Analysis: {total_matches} total match(es) found\n\n"
        message += f"📊 Summary:\n"
        message += f"Files analyzed: {len(results)}\n"
        message += f"Total matches: {total_matches} (in {files_with_matches} files)\n"

        
        message += f"Search text: '{search_text[:40]}{'...' if len(search_text) > 40 else ''}'\n"
        message += f"Replace text: '{replace_text[:40]}{'...' if len(replace_text) > 40 else ''}'\n"
        message += "=" * 77 + "\n\n"
        
        # Show detailed file results
        for i, result in enumerate(results, 1):
            message += f"📁 FILE {i}: {result['filename']}\n"
            
            if result['total'] > 0:
                message += f"   ✅ Total matches: {result['total']}\n"
            else:
                message += f"   ✅ No matches found\n"

            
            for detail in result['details']:
                message += f"   {detail}\n"
            message += "\n" + "=" * 77 + "\n\n"
        
        # Recommendations
        if total_matches > 0:
            message += "💡 Replacement recommendations:\n"
            message += "Use 'Advanced Replace' for complete coverage with hyperlink preservation\n"
            message += "Or use 'Standard Replace' for faster processing (main content only)\n"
        else:
            message += "💡 Tip: No matches found. Try different search terms or check spelling."
        
        self.show_scrollable_results(message, "Advanced Preview results")

    def replace_in_paragraph_advanced(self, paragraph, search_text, replace_text, case_sensitive=False, use_regex=False, whole_word=False):
        """Replace text in a paragraph that may span multiple runs"""
        # Strip invisible formatting characters from runs before matching
        for run in paragraph.runs:
            cleaned = self._strip_invisible_chars(run.text)
            if cleaned != run.text:
                run.text = cleaned
        
        paragraph_text = paragraph.text
        
        # If whole_word and not regex, use \b boundaries internally
        effective_regex = use_regex
        effective_search = search_text
        if whole_word and not use_regex:
            effective_search = r'\b' + re.escape(search_text) + r'\b'
            effective_regex = True
        
        # Count matches
        if effective_regex:
            try:
                flags = 0 if case_sensitive else re.IGNORECASE
                replacements = len(re.findall(effective_search, paragraph_text, flags))
            except re.error:
                return 0
            if replacements == 0:
                return 0
        else:
            if case_sensitive:
                if search_text not in paragraph_text:
                    return 0
                replacements = paragraph_text.count(search_text)
            else:
                if search_text.lower() not in paragraph_text.lower():
                    return 0
                replacements = paragraph_text.lower().count(search_text.lower())
        
        # If the search text is contained within a single run, use simple replacement
        for run in paragraph.runs:
            run_text = run.text
            if effective_regex:
                try:
                    flags = 0 if case_sensitive else re.IGNORECASE
                    if re.search(effective_search, run_text, flags):
                        run.text = re.sub(effective_search, replace_text, run_text, flags=flags)
                        return replacements
                except re.error:
                    return 0
            elif case_sensitive:
                if search_text in run_text:
                    run.text = run_text.replace(search_text, replace_text)
                    return replacements
            else:
                if search_text.lower() in run_text.lower():
                    pattern = re.escape(search_text)
                    run.text = re.sub(pattern, replace_text, run_text, flags=re.IGNORECASE)
                    return replacements
        
        # If we get here, the text spans multiple runs - rebuild the paragraph
        if effective_regex:
            try:
                flags = 0 if case_sensitive else re.IGNORECASE
                new_paragraph_text = re.sub(effective_search, replace_text, paragraph_text, flags=flags)
            except re.error:
                return 0
        elif case_sensitive:
            new_paragraph_text = paragraph_text.replace(search_text, replace_text)
        else:
            pattern = re.escape(search_text)
            new_paragraph_text = re.sub(pattern, replace_text, paragraph_text, flags=re.IGNORECASE)
        
        # Clear all runs and add new text to first run
        for run in paragraph.runs:
            run.text = ""
        
        if paragraph.runs:
            paragraph.runs[0].text = new_paragraph_text
        else:
            paragraph.text = new_paragraph_text
        
        return replacements
#########END OF PART 3A######






#########PART 3B######
    def _replace_in_table(self, table, search_text, replace_text, case_sensitive, use_regex=False, whole_word=False):
        """Replace text in a table, recursing into nested tables"""
        count = 0
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count += self.replace_in_paragraph_advanced(paragraph, search_text, replace_text, case_sensitive, use_regex, whole_word)
                for nested_table in cell.tables:
                    count += self._replace_in_table(nested_table, search_text, replace_text, case_sensitive, use_regex, whole_word)
        return count

    def replace_text_in_documents(self, close_after=False):
        """Perform the actual text replacement across all files - Enhanced results"""
        if not self.file_paths:
            messagebox.showwarning("Warning", "Please add at least one Word document.")
            return False
            
        search_for = self.get_processed_search_text()    # CHANGED: Use processed text
        replace_with = self.get_processed_replace_text() # CHANGED: Use processed text
        
        if not search_for:
            messagebox.showwarning("Warning", "Please enter text to search for.")
            return False
        
        use_regex = self.regex_var.get()
        
        # Validate regex pattern
        if use_regex:
            try:
                re.compile(search_for)
            except re.error as e:
                messagebox.showerror("Invalid Regex", f"Invalid regex pattern:\n{str(e)}")
                return False
        
        # Confirm action
        file_count = len(self.file_paths)
        case_sensitive = self.case_sensitive_var.get()
        case_info = " (case sensitive)" if case_sensitive else " (case insensitive)"
        regex_info = " [REGEX]" if use_regex else ""
        action_text = "replace all occurrences and close the application" if close_after else "replace all occurrences"
        if not messagebox.askyesno("Confirm", 
                                  f"This will {action_text} in {file_count} file(s){case_info}{regex_info}.\n\nContinue?"):
            return False
        
        try:
            self.status_label.config(text="Processing files...", fg="orange")
            self.progress_frame.pack(fill=tk.X, pady=(10, 0))
            self.root.update()
            
            total_replacements = 0
            successful_files = 0
            backup_files = []
            all_results = []  # NEW: Store detailed results
            
            for i, file_path in enumerate(self.file_paths):
                # Update progress with percentage
                progress_percent = int((i / file_count) * 100)
                self.progress_label.config(text=f"Processing {i+1}/{file_count}: {os.path.basename(file_path)}")
                self.progress_percent_label.config(text=f"Progress: {progress_percent}%")
                self.root.update()
                
                try:
                    # Create backup if requested
                    if self.create_backup_var.get():
                        backup_path = file_path + ".backup"
                        import shutil
                        shutil.copy2(file_path, backup_path)
                        backup_files.append(backup_path)
                    
                    # Process the document
                    doc = Document(file_path)
                    file_replacements = 0
                    whole_word = self.whole_word_var.get()
                    
                    # Replace in paragraphs
                    for paragraph in doc.paragraphs:
                        file_replacements += self.replace_in_paragraph_advanced(paragraph, search_for, replace_with, case_sensitive, use_regex, whole_word)
                    
                    # Replace in tables (including nested tables)
                    for table in doc.tables:
                        file_replacements += self._replace_in_table(table, search_for, replace_with, case_sensitive, use_regex, whole_word)
                    
                    # Save the document
                    doc.save(file_path)
                    
                    # Store detailed results
                    all_results.append({
                        'filename': os.path.basename(file_path),
                        'total': file_replacements,
                        'details': [f"  📄 Main content replacements: {file_replacements}"] if file_replacements > 0 else ["  ✅ No replacements needed"]
                    })
                    
                    total_replacements += file_replacements
                    successful_files += 1
                    
                except Exception as e:
                    all_results.append({
                        'filename': os.path.basename(file_path),
                        'total': 0,
                        'details': [f"  ❌ Error: {str(e)}"]
                    })
            
            # Final progress update
            self.progress_percent_label.config(text="Progress: 100%")
            self.root.update()
            
            self.progress_frame.pack_forget()
            self.status_label.config(text="Replacement completed!", fg="green")
            
            # Refresh text cache so live counter reflects the updated files
            self._refresh_text_cache()
            self._schedule_live_count()
            
            # Show enhanced results instead of simple messagebox
            self.show_replace_all_results(all_results, search_for, replace_with, total_replacements, 
                                         successful_files, backup_files, close_after)
            
            return True
            
        except Exception as e:
            self.status_label.config(text="Error occurred", fg="red")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            return False

    def show_replace_all_results(self, results, search_text, replace_text, total_replacements, successful_files, backup_files, close_after):
        """Show Replace results in scrollable window"""
        case_sensitive = self.case_sensitive_var.get()
        case_info = " (case sensitive)" if case_sensitive else " (case insensitive)"
        
        message = "-" * 46 + "\n"
        message += f"✅ Standard Replace results{case_info}\n"
        message += "-" * 46 + "\n\n" 
        message += f"🎉 Completed: {total_replacements} replacement(s) in main document content\n\n"
        message += f"📊 Summary:\n"
        message += f"Files processed: {successful_files}/{len(self.file_paths)}\n"
        message += f"Total replacements: {total_replacements}\n"
        message += f"Search text: '{search_text[:40]}{'...' if len(search_text) > 40 else ''}'\n"
        message += f"Replace text: '{replace_text[:40]}{'...' if len(replace_text) > 40 else ''}'\n"
        message += f"Coverage: Main text and tables\n"
        
        if backup_files:
            message += f"Backup files created: {len(backup_files)}\n"
        
        message += "=" * 77 + "\n\n"
        
        # Show detailed file results
        for i, result in enumerate(results, 1):
            message += f"📁 FILE {i}: {result['filename']}\n"
            if result['total'] > 0:
                message += f"   ✅ Replacements made: {result['total']}\n"
            else:
                message += f"   ✅ No replacements needed\n"
            
            for detail in result['details']:
                message += f"   {detail}\n"
            message += "\n" + "=" * 77 + "\n\n"
        
        if close_after:
            message += "🚪 The application will close after you click OK.\n\n"
        
        # Show results and handle close_after
        if close_after:
            # For close_after, show results then close
            result_window = tk.Toplevel(self.root)
            result_window.title("Standard Replace results")
            result_window.geometry("600x500")
            result_window.resizable(True, True)
            
            text_widget = scrolledtext.ScrolledText(result_window, wrap=tk.WORD, 
                                                   font=("Consolas", 9), padx=10, pady=10)
            text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            text_widget.insert("1.0", message)
            text_widget.config(state=tk.DISABLED)
            
            def close_and_exit():
                result_window.destroy()
                self.root.destroy()
            
            close_btn = tk.Button(result_window, text="Close application", 
                                 command=close_and_exit,
                                 bg="#333333", fg="white", font=("Arial", 10, "bold"))
            close_btn.pack(pady=(0, 10))
            
            result_window.transient(self.root)
            result_window.grab_set()
        else:
            # Normal scrollable results
            self.show_scrollable_results(message, "Standard Replace results")

    # Advanced Replace with Word COM Automation
    def advanced_replace_with_vba(self):
        """Advanced replace using Word COM automation for ALL document areas"""
        if not self.file_paths:
            messagebox.showwarning("Warning", "Please add at least one Word document.")
            return
        
        if not HAS_WIN32COM:
            messagebox.showerror("Error", 
                "Advanced replace requires the 'pywin32' package.\n\n"
                "Install it with: pip install pywin32")
            return
        
        if self.regex_var.get():
            messagebox.showinfo("Regex Not Supported", 
                "Regex is only supported with Standard Replace.\n\n"
                "Advanced Replace uses Word's native Find & Replace engine\n"
                "which does not support Python regex syntax.\n\n"
                "Please uncheck 'Regex' or use Standard Replace instead.")
            return
            
        search_for = self.get_processed_search_text()
        replace_with = self.get_processed_replace_text()
        
        if not search_for:
            messagebox.showwarning("Warning", "Please enter text to search for.")
            return
        
        # Confirm action
        file_count = len(self.file_paths)
        case_sensitive = self.case_sensitive_var.get()
        case_info = " (case sensitive)" if case_sensitive else " (case insensitive)"
        if not messagebox.askyesno("Confirm Advanced Replace", 
                                  f"This will replace text in all document areas in {file_count} file(s){case_info}.\n\n" +
                                  f"Areas covered: Main content, shapes, headers, footers, footnotes, endnotes, form fields, hyperlinks\n\n" +
                                  f"Search: '{search_for[:30]}{'...' if len(search_for) > 30 else ''}'\n" +
                                  f"Replace: '{replace_with[:30]}{'...' if len(replace_with) > 30 else ''}'\n\n" +
                                  "Continue?"):
            return
        
        word_app = None
        try:
            self.status_label.config(text="Processing advanced replacements...", fg="orange")
            self.progress_frame.pack(fill=tk.X, pady=(10, 0))
            self.root.update()
            
            all_results = []
            total_advanced_replacements = 0
            successful_files = 0
            errors = []
            
            # Open Word once for all files
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False
            word_app.ScreenUpdating = False
            
            for i, file_path in enumerate(self.file_paths):
                # Update progress
                progress_percent = int((i / file_count) * 100)
                self.progress_label.config(text=f"Processing {i+1}/{file_count}: {os.path.basename(file_path)}")
                self.progress_percent_label.config(text=f"Progress: {progress_percent}%")
                self.root.update()
                
                doc = None
                try:
                    # Create backup if requested
                    if self.create_backup_var.get():
                        backup_path = file_path + ".backup"
                        import shutil
                        shutil.copy2(file_path, backup_path)
                    
                    full_path = os.path.abspath(file_path)
                    doc = word_app.Documents.Open(full_path)
                    
                    file_replacements = 0
                    details = []
                    
                    # Pre-cache shape ranges for O(n) hyperlink-in-shape check
                    shape_ranges = self._build_shape_ranges(doc)
                    whole_word = self.whole_word_var.get()
                    
                    # Replace in text boxes / shapes via StoryRanges (wdTextFrameStory = 5)
                    # This catches ALL text boxes regardless of how they were inserted
                    shape_count = 0
                    try:
                        story = doc.StoryRanges(5)  # wdTextFrameStory
                        while story:
                            if len(story.Text) > 1:
                                shape_count += self._find_replace_count(
                                    story, search_for, replace_with, case_sensitive, whole_word)
                            try:
                                story = story.NextStoryRange
                            except Exception:
                                break
                    except Exception:
                        pass
                    if shape_count > 0:
                        details.append(f"  📦 Text boxes: {shape_count} replacement(s)")
                        file_replacements += shape_count
                    
                    # Replace in headers
                    header_count = 0
                    for section in doc.Sections:
                        for idx in range(1, 4):
                            try:
                                if section.Headers(idx).Exists:
                                    if len(section.Headers(idx).Range.Text) > 1:
                                        header_count += self._find_replace_count(
                                            section.Headers(idx).Range, search_for, replace_with, case_sensitive, whole_word)
                            except Exception:
                                pass
                    if header_count > 0:
                        details.append(f"  📄 Headers: {header_count} replacement(s)")
                        file_replacements += header_count
                    
                    # Replace in footers
                    footer_count = 0
                    for section in doc.Sections:
                        for idx in range(1, 4):
                            try:
                                if section.Footers(idx).Exists:
                                    if len(section.Footers(idx).Range.Text) > 1:
                                        footer_count += self._find_replace_count(
                                            section.Footers(idx).Range, search_for, replace_with, case_sensitive, whole_word)
                            except Exception:
                                pass
                    if footer_count > 0:
                        details.append(f"  📄 Footers: {footer_count} replacement(s)")
                        file_replacements += footer_count
                    
                    # Replace in footnotes
                    footnote_count = 0
                    for footnote in doc.Footnotes:
                        try:
                            if len(footnote.Range.Text) > 1:
                                footnote_count += self._find_replace_count(
                                    footnote.Range, search_for, replace_with, case_sensitive, whole_word)
                        except Exception:
                            pass
                    if footnote_count > 0:
                        details.append(f"  📝 Footnotes: {footnote_count} replacement(s)")
                        file_replacements += footnote_count
                    
                    # Replace in endnotes
                    endnote_count = 0
                    for endnote in doc.Endnotes:
                        try:
                            if len(endnote.Range.Text) > 1:
                                endnote_count += self._find_replace_count(
                                    endnote.Range, search_for, replace_with, case_sensitive, whole_word)
                        except Exception:
                            pass
                    if endnote_count > 0:
                        details.append(f"  📝 Endnotes: {endnote_count} replacement(s)")
                        file_replacements += endnote_count
                    
                    # Replace in form fields (no Find.Execute — direct text manipulation)
                    form_field_count = 0
                    for field in doc.FormFields:
                        try:
                            if field.Type == 70:  # wdFieldFormTextInput
                                original_text = field.Result
                                if len(original_text) > 0:
                                    # Strip invisible chars for matching AND replacement
                                    clean_text = self._strip_invisible_chars(original_text)
                                    occ = self.count_occurrences(clean_text, search_for, case_sensitive, False, whole_word)
                                    if occ > 0:
                                        if whole_word:
                                            pattern = r'\b' + re.escape(search_for) + r'\b'
                                            flags = 0 if case_sensitive else re.IGNORECASE
                                            new_text = re.sub(pattern, replace_with, clean_text, flags=flags)
                                        elif case_sensitive:
                                            new_text = clean_text.replace(search_for, replace_with)
                                        else:
                                            pattern = re.escape(search_for)
                                            new_text = re.sub(pattern, replace_with, clean_text, flags=re.IGNORECASE)
                                        field.Result = new_text
                                        form_field_count += occ
                        except Exception:
                            pass
                    if form_field_count > 0:
                        details.append(f"  📋 Form fields: {form_field_count} replacement(s)")
                        file_replacements += form_field_count
                    
                    # Replace in main document content using iterative Find (Option A)
                    # (this covers paragraphs, tables, and inline hyperlink text)
                    main_content_count = self._find_replace_count(
                        doc.Content, search_for, replace_with, case_sensitive, whole_word)
                    
                    # Count hyperlink replacements that were part of main content
                    # (for accurate per-area reporting, not for additional replacement)
                    hyperlink_count = 0
                    for hyperlink in doc.Hyperlinks:
                        try:
                            hl_range = hyperlink.Range
                            if hl_range is None:
                                continue
                            if not self._is_in_shape_ranges(hl_range.Start, hl_range.End, shape_ranges):
                                display_text = hl_range.Text
                                if display_text and len(display_text) > 0:
                                    occ = self.count_occurrences(display_text, replace_with, case_sensitive, False, whole_word) if replace_with else 0
                                    if occ > 0:
                                        hyperlink_count += occ
                        except Exception:
                            pass
                    
                    # Report main content minus hyperlink portion to avoid double-counting
                    non_hyperlink_main = max(0, main_content_count - hyperlink_count)
                    if non_hyperlink_main > 0:
                        details.append(f"  📄 Main content: {non_hyperlink_main} replacement(s)")
                        file_replacements += non_hyperlink_main
                    if hyperlink_count > 0:
                        details.append(f"  🔗 Hyperlinks: {hyperlink_count} replacement(s)")
                        file_replacements += hyperlink_count
                    
                    # Save and close document
                    doc.Save()
                    doc.Close()
                    doc = None
                    
                    all_results.append({
                        'filename': os.path.basename(file_path),
                        'total': file_replacements,
                        'details': details if details else ["  ✅ No advanced matches found"]
                    })
                    
                    total_advanced_replacements += file_replacements
                    successful_files += 1
                    
                except Exception as e:
                    if doc:
                        try:
                            doc.Close(False)
                        except Exception:
                            pass
                        doc = None
                    all_results.append({
                        'filename': os.path.basename(file_path),
                        'total': 0,
                        'details': [f"  ❌ Error: {str(e)}"]
                    })
                    errors.append(f"{os.path.basename(file_path)}: {str(e)}")
            
            # Final progress update
            self.progress_percent_label.config(text="Progress: 100%")
            self.root.update()
            
            self.progress_frame.pack_forget()
            self.status_label.config(text="Advanced replacement completed!", fg="green")
            
            # Refresh text cache so live counter reflects the updated files
            self._refresh_text_cache()
            self._schedule_live_count()
            
            # Show results in scrollable window
            self.show_advanced_replace_results(all_results, search_for, replace_with, total_advanced_replacements, successful_files, errors)
            
        except Exception as e:
            self.status_label.config(text="Error occurred", fg="red")
            messagebox.showerror("Error", f"Advanced replace error: {str(e)}")
        finally:
            try:
                if word_app:
                    word_app.ScreenUpdating = True
                    word_app.Quit()
            except Exception:
                pass

    def show_advanced_replace_results(self, results, search_text, replace_text, total_replacements, successful_files, errors):
        """Show advanced replace results in scrollable window"""
        case_sensitive = self.case_sensitive_var.get()
        case_info = " (case sensitive)" if case_sensitive else " (case insensitive)"
        
        message = "-" * 46 + "\n"
        message += f"🔧 Advanced Replace results{case_info}\n"
        message += "-" * 46 + "\n\n"
        message += f"✅ Completed: {total_replacements} replacement(s) in ALL document areas\n\n"
        message += f"📊 Summary:\n"
        message += f"Files processed: {successful_files}/{len(self.file_paths)}\n"
        message += f"Total advanced replacements: {total_replacements}\n"
        message += f"Search text: '{search_text[:40]}{'...' if len(search_text) > 40 else ''}'\n"
        message += f"Replace text: '{replace_text[:40]}{'...' if len(replace_text) > 40 else ''}'\n"
        message += f"Coverage: Complete document (main content + all advanced areas)\n"
        message += "=" * 77 + "\n\n"
        
        # Show detailed file results
        for i, result in enumerate(results, 1):
            message += f"📁 FILE {i}: {result['filename']}\n"
            if result['total'] > 0:
                message += f"   ✅ Advanced replacements: {result['total']}\n"
            else:
                message += f"   ✅ No advanced matches found\n"
            
            for detail in result['details']:
                message += f"   {detail}\n"
            message += "\n" + "=" * 77 + "\n\n"
        
        if errors:
            message += f"❌ Errors:\n"
            for error in errors:
                message += f"• {error}\n"
            message += "\n"
        
        self.show_scrollable_results(message, "Advanced Replace results")
    
    def show_scrollable_results(self, message, title):
        """Show results in a scrollable window for long messages"""
        result_window = tk.Toplevel(self.root)
        result_window.title(title)
        result_window.geometry("600x500")
        result_window.resizable(True, True)
        
        # Create scrolled text widget
        text_widget = scrolledtext.ScrolledText(result_window, wrap=tk.WORD, 
                                               font=("Consolas", 9), padx=10, pady=10)
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Insert the message
        text_widget.insert("1.0", message)
        text_widget.config(state=tk.DISABLED)  # Make it read-only
        
        # Button frame for two buttons
        button_frame = tk.Frame(result_window)
        button_frame.pack(pady=(0, 10))
        
        # OK button (just closes result window)
        ok_btn = tk.Button(button_frame, text="OK", 
                          command=result_window.destroy,
                          bg="#4caf50", fg="white", font=("Arial", 10))
        ok_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Close application button (closes result window AND main app)
        def close_application():
            result_window.destroy()
            self.root.destroy()
        
        close_app_btn = tk.Button(button_frame, text="Close application", 
                                 command=close_application,
                                 bg="#333333", fg="white", font=("Arial", 10))
        close_app_btn.pack(side=tk.LEFT)
        
        # Center the window
        result_window.transient(self.root)
        result_window.grab_set()
#########END OF PART 3B######




#########PART 4 (FINAL)######
    def run(self):
        self.root.mainloop()

def main():
    # Get the initial file from command line (if any)
    initial_file = sys.argv[1] if len(sys.argv) > 1 else None
    
    # Validate the initial file
    if initial_file and not os.path.exists(initial_file):
        initial_file = None
    
    if initial_file and not initial_file.lower().endswith(('.docx', '.doc', '.docm')):
        initial_file = None
    
    try:
        app = WordTextReplacerSingle(initial_file)
        app.run()
        
    except Exception as e:
        messagebox.showerror("Error", f"Failed to start application: {str(e)}")

if __name__ == "__main__":
    main()

#########END OF PART 4 (FINAL)######
